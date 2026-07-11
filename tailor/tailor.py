#!/usr/bin/env python3
# Reverse-Recruiter Job Search System - tailor
# Copyright (C) 2026 Jennifer McKinney
# This program is free software: you can redistribute it and/or modify it under
# the terms of the GNU General Public License v3.0 as published by the Free
# Software Foundation. Distributed WITHOUT ANY WARRANTY. See
# <https://www.gnu.org/licenses/>. Commercial licensing available from the author.
"""Generate a tailored cover letter and a simple resume from your config and a role spec.

Usage:
    python3 tailor.py --config ../config/config.json --role role.example.json --out ./output

Inputs:
    config.json  - your profile (see config.example.json). Uses the "operator" block,
                   plus optional "summary" (str) and "resume" {skills[], experience[]}.
    role.json    - the target role: {"company","title","hiring_manager"?,"highlights"?[]}

Outputs (in --out):
    CoverLetter_<Company>.docx
    Resume_<YourName>.docx
"""
__version__ = "1.0.0"

import argparse
import json
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    sys.exit("python-docx is required. Install it with:  pip install python-docx")


def load_json(path: str) -> dict:
    p = Path(path)
    if not p.is_file():
        sys.exit(f"File not found: {path}")
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        sys.exit(f"Invalid JSON in {path}: {exc}")


def _para(doc: Document, text: str, size: int = 11, bold: bool = False, space_after: int = 8) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(space_after)


def build_cover_letter(cfg: dict, role: dict, outdir: Path) -> Path:
    op = cfg.get("operator", {})
    company = role.get("company", "the company")
    title = role.get("title", "the role")
    greeting_name = role.get("hiring_manager") or "Hiring Team"
    highlights = role.get("highlights") or []
    summary = cfg.get("summary", "").strip()

    doc = Document()
    _para(doc, date.today().strftime("%B %d, %Y"), space_after=12)
    _para(doc, f"Dear {greeting_name},", space_after=10)
    _para(doc, f"I'm writing to express my interest in the {title} role at {company}.", space_after=10)
    if summary:
        _para(doc, summary, space_after=10)
    if highlights:
        _para(doc, f"A few reasons I'd be a strong fit for {company}:", bold=True, space_after=6)
        for h in highlights:
            b = doc.add_paragraph(style="List Bullet")
            b.add_run(str(h)).font.size = Pt(11)
    _para(
        doc,
        f"I'd welcome the chance to discuss how I can contribute to {company}. "
        "Thank you for your time and consideration.",
        space_after=12,
    )
    _para(doc, "Sincerely,", space_after=2)
    _para(doc, op.get("name", "Your Name"), bold=True, space_after=2)
    contact = " | ".join(x for x in [op.get("email", ""), op.get("phone", ""), op.get("linkedin", "")] if x)
    if contact:
        _para(doc, contact, size=10)

    safe = "".join(ch for ch in company if ch.isalnum() or ch in " -_").strip().replace(" ", "_")
    out = outdir / f"CoverLetter_{safe or 'Company'}.docx"
    doc.save(out)
    return out


def build_resume(cfg: dict, outdir: Path) -> Path:
    op = cfg.get("operator", {})
    resume = cfg.get("resume", {}) or {}

    doc = Document()
    _para(doc, op.get("name", "Your Name"), size=18, bold=True, space_after=2)
    contact = " | ".join(x for x in [op.get("email", ""), op.get("phone", ""), op.get("linkedin", ""), op.get("location_line", "")] if x)
    if contact:
        _para(doc, contact, size=10, space_after=10)
    if cfg.get("summary"):
        _para(doc, "SUMMARY", bold=True, space_after=4)
        _para(doc, cfg["summary"], space_after=10)
    if resume.get("skills"):
        _para(doc, "SKILLS", bold=True, space_after=4)
        _para(doc, ", ".join(str(s) for s in resume["skills"]), space_after=10)
    if resume.get("experience"):
        _para(doc, "EXPERIENCE", bold=True, space_after=4)
        for job in resume["experience"]:
            line = f"{job.get('title','')} - {job.get('company','')} ({job.get('dates','')})".strip(" -")
            _para(doc, line, bold=True, space_after=2)
            for bullet in job.get("bullets", []):
                b = doc.add_paragraph(style="List Bullet")
                b.add_run(str(bullet)).font.size = Pt(11)
    else:
        _para(doc, "Add a \"resume\" block to config.json to populate experience.", size=10)

    name_safe = "".join(ch for ch in op.get("name", "Resume") if ch.isalnum()) or "Resume"
    out = outdir / f"Resume_{name_safe}.docx"
    doc.save(out)
    return out


def main() -> None:
    ap = argparse.ArgumentParser(description="Tailor a cover letter and resume from config + role.")
    ap.add_argument("--config", required=True, help="Path to config.json")
    ap.add_argument("--role", required=True, help="Path to role.json (company, title, highlights)")
    ap.add_argument("--out", default="./output", help="Output directory (default: ./output)")
    args = ap.parse_args()

    cfg = load_json(args.config)
    role = load_json(args.role)
    outdir = Path(args.out)
    outdir.mkdir(parents=True, exist_ok=True)

    cover = build_cover_letter(cfg, role, outdir)
    resume = build_resume(cfg, outdir)
    print(f"Created: {cover}")
    print(f"Created: {resume}")


if __name__ == "__main__":
    main()
